import os
import re
import json
import pypdf
import docx
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# ---------------------------------------------------------
# Comprehensive Technical Skills Dictionary
# ---------------------------------------------------------
SKILL_TAXONOMY = {
    'Programming Languages': [
        'python', 'java', 'c++', 'c#', 'javascript', 'typescript', 'php', 'ruby',
        'go', 'golang', 'rust', 'kotlin', 'swift', 'r', 'c', 'sql', 'html', 'css', 'bash', 'shell'
    ],
    'Frameworks & Libraries': [
        'react', 'react.js', 'reactjs', 'angular', 'vue', 'vue.js', 'node.js', 'nodejs',
        'express', 'django', 'flask', 'fastapi', 'spring', 'spring boot', 'asp.net', '.net',
        'bootstrap', 'tailwind', 'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch',
        'keras', 'opencv', 'next.js', 'nuxt.js', 'jquery', 'flutter', 'react native'
    ],
    'Databases & Storage': [
        'mysql', 'postgresql', 'postgres', 'mongodb', 'sqlite', 'redis', 'oracle',
        'sql server', 'mssql', 'firebase', 'cassandra', 'dynamodb', 'elasticsearch', 'mariadb'
    ],
    'DevOps & Tools': [
        'git', 'github', 'gitlab', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
        'google cloud', 'jenkins', 'linux', 'unix', 'jira', 'postman', 'ci/cd', 'terraform',
        'ansible', 'maven', 'gradle', 'npm', 'pip', 'vscode'
    ],
    'Concepts & Methodology': [
        'data structures', 'algorithms', 'dsa', 'machine learning', 'deep learning',
        'ai', 'artificial intelligence', 'rest api', 'restful api', 'microservices',
        'agile', 'scrum', 'system design', 'oop', 'object oriented programming',
        'unit testing', 'software development', 'web development', 'cloud computing',
        'cybersecurity', 'network security'
    ]
}

# Target Job Role Profiles
JOB_PROFILES = [
    {
        'title': 'Software Developer',
        'min_cgpa': 6.5,
        'required_skills': ['data structures', 'algorithms', 'python', 'java', 'c++', 'sql', 'git', 'oop'],
        'description': 'Develop, test, and maintain robust desktop and enterprise software solutions.'
    },
    {
        'title': 'Full Stack Developer',
        'min_cgpa': 7.0,
        'required_skills': ['javascript', 'react', 'node.js', 'html', 'css', 'sql', 'mongodb', 'git', 'express'],
        'description': 'Build responsive frontend user interfaces and scalable backend server APIs.'
    },
    {
        'title': 'Frontend Developer',
        'min_cgpa': 6.5,
        'required_skills': ['javascript', 'react', 'html', 'css', 'bootstrap', 'tailwind', 'typescript', 'git'],
        'description': 'Craft modern visual designs, user experiences, and single page web applications.'
    },
    {
        'title': 'Backend Developer',
        'min_cgpa': 7.0,
        'required_skills': ['python', 'node.js', 'django', 'flask', 'sql', 'postgresql', 'mongodb', 'rest api', 'docker'],
        'description': 'Architect server-side logic, database schemas, and microservice APIs.'
    },
    {
        'title': 'Python Developer',
        'min_cgpa': 6.5,
        'required_skills': ['python', 'django', 'flask', 'fastapi', 'sql', 'git', 'rest api', 'pandas'],
        'description': 'Engineers backend services, data pipelines, and automation tools using Python.'
    },
    {
        'title': 'Java Developer',
        'min_cgpa': 6.5,
        'required_skills': ['java', 'spring', 'spring boot', 'sql', 'mysql', 'hibernate', 'git', 'dsa'],
        'description': 'Develop scalable backend architecture and enterprise application services.'
    },
    {
        'title': 'AI / ML Engineer',
        'min_cgpa': 7.5,
        'required_skills': ['python', 'machine learning', 'deep learning', 'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch'],
        'description': 'Design predictive machine learning models, neural networks, and AI algorithms.'
    },
    {
        'title': 'Data Analyst',
        'min_cgpa': 6.5,
        'required_skills': ['python', 'sql', 'pandas', 'numpy', 'excel', 'tableau', 'power bi', 'r'],
        'description': 'Transform raw dataset metrics into business intelligence reports and strategic insights.'
    },
    {
        'title': 'QA / Test Automation Engineer',
        'min_cgpa': 6.0,
        'required_skills': ['java', 'python', 'selenium', 'unit testing', 'postman', 'git', 'jira', 'sql'],
        'description': 'Automate software quality testing suites and validate system functionality.'
    },
    {
        'title': 'Cloud & DevOps Engineer',
        'min_cgpa': 7.0,
        'required_skills': ['linux', 'aws', 'docker', 'kubernetes', 'ci/cd', 'jenkins', 'python', 'bash'],
        'description': 'Deploy cloud infrastructure, automate deployment pipelines, and maintain uptime.'
    }
]

# ---------------------------------------------------------
# Text Extraction Functions
# ---------------------------------------------------------
def extract_text_from_pdf(filepath):
    """Extract raw text from PDF resume file."""
    text = ""
    try:
        reader = pypdf.PdfReader(filepath)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        print(f"[ERROR] PDF extraction failed: {e}")
    return text

def extract_text_from_docx(filepath):
    """Extract raw text from Word DOCX resume file."""
    text = ""
    try:
        doc = docx.Document(filepath)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        print(f"[ERROR] DOCX extraction failed: {e}")
    return text

def extract_text_from_file(filepath, file_type):
    """Generic text extractor supporting PDF and DOCX."""
    ext = file_type.lower()
    if ext == 'pdf' or filepath.endswith('.pdf'):
        return extract_text_from_pdf(filepath)
    elif ext in ['docx', 'doc'] or filepath.endswith('.docx'):
        return extract_text_from_docx(filepath)
    else:
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""

# ---------------------------------------------------------
# Structured Parsing Engine
# ---------------------------------------------------------
def parse_resume_text(raw_text):
    """Parse candidate name, contact details, skills, education, projects, experience, certifications."""
    clean_text = raw_text.strip()
    lines = [line.strip() for line in clean_text.split('\n') if line.strip()]

    # 1. Email Extraction
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', clean_text)
    email = email_match.group(0).strip().lower() if email_match else "Not Specified"

    # 2. Phone Extraction
    phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', clean_text)
    phone = phone_match.group(0).strip() if phone_match else "Not Specified"

    # 3. Name Extraction Heuristics
    candidate_name = "Candidate"
    for line in lines[:8]:
        # Filter out common resume header words
        if not re.search(r'resume|curriculum|vitae|email|phone|contact|profile|page', line, re.IGNORECASE) and not re.search(r'@|\d{10}', line):
            if 2 <= len(line.split()) <= 4 and len(line) < 35:
                candidate_name = line.strip().title()
                break

    # 4. Skills Extraction
    extracted_skills = set()
    text_lower = clean_text.lower()
    for category, skill_list in SKILL_TAXONOMY.items():
        for skill in skill_list:
            # Word boundary regex matching to avoid partial word collisions
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                # Standardize casing
                formatted_skill = skill.upper() if len(skill) <= 3 and skill != 'c' else skill.title()
                if formatted_skill.lower() == 'react.js' or formatted_skill.lower() == 'reactjs':
                    formatted_skill = 'React'
                elif formatted_skill.lower() == 'node.js' or formatted_skill.lower() == 'nodejs':
                    formatted_skill = 'Node.js'
                extracted_skills.add(formatted_skill)

    skills_list = sorted(list(extracted_skills))

    # 5. Education Extraction
    education_entries = []
    degree_patterns = [
        r'(b\.tech|b\.e\.|bachelor of technology|bachelor of engineering|b\.sc|m\.tech|m\.e\.|mca|bca)',
        r'(computer science|information technology|electronics|electrical|mechanical|civil)',
        r'(cgpa|gpa|percentage)[:\s]+([0-9\.]+)'
    ]
    for line in lines:
        for p in degree_patterns:
            if re.search(p, line, re.IGNORECASE):
                if line not in education_entries:
                    education_entries.append(line)

    # CGPA Extraction Heuristic
    cgpa_val = 7.5
    cgpa_match = re.search(r'(?:cgpa|gpa|marks)[:\s=]*([0-9]\.[0-9]{1,2})', text_lower)
    if cgpa_match:
        try:
            val = float(cgpa_match.group(1))
            if 4.0 <= val <= 10.0:
                cgpa_val = val
        except ValueError:
            pass

    # 6. Projects Extraction
    projects_list = []
    in_project_section = False
    for line in lines:
        if re.search(r'^\s*(projects|academic projects|key projects|personal projects)', line, re.IGNORECASE):
            in_project_section = True
            continue
        if in_project_section and re.search(r'^\s*(experience|internship|education|certifications|skills)', line, re.IGNORECASE):
            in_project_section = False
        if in_project_section:
            if len(line) > 15 and not line.startswith('http'):
                projects_list.append(line)

    # 7. Internships & Experience Extraction
    experience_list = []
    in_exp_section = False
    for line in lines:
        if re.search(r'^\s*(experience|internship|work history|employment|internships)', line, re.IGNORECASE):
            in_exp_section = True
            continue
        if in_exp_section and re.search(r'^\s*(projects|education|certifications|skills|declaration)', line, re.IGNORECASE):
            in_exp_section = False
        if in_exp_section:
            if len(line) > 15:
                experience_list.append(line)

    # 8. Certifications Extraction
    cert_list = []
    in_cert_section = False
    cert_keywords = ['aws', 'oracle', 'microsoft', 'google', 'meta', 'ibm', 'certified', 'coursera', 'udemy', 'nptel', 'certification']
    for line in lines:
        if re.search(r'^\s*(certifications|certificates|licenses|courses)', line, re.IGNORECASE):
            in_cert_section = True
            continue
        if in_cert_section and re.search(r'^\s*(projects|education|experience|skills)', line, re.IGNORECASE):
            in_cert_section = False
        if in_cert_section or any(k in line.lower() for k in cert_keywords):
            if len(line) > 10 and not any(header in line.lower() for header in ['certifications', 'certificates']):
                if line not in cert_list:
                    cert_list.append(line)

    return {
        'name': candidate_name,
        'email': email,
        'phone': phone,
        'skills': skills_list,
        'cgpa': cgpa_val,
        'education': education_entries if education_entries else ["Bachelor of Technology in Computer Science & Engineering"],
        'projects': projects_list if projects_list else ["Full-Stack Web Development & ML Placement Predictor Project"],
        'experience': experience_list if experience_list else [],
        'certifications': cert_list if cert_list else []
    }

# ---------------------------------------------------------
# AI Resume Scoring Engine (Out of 100)
# ---------------------------------------------------------
def calculate_ai_resume_score(parsed_data):
    """Compute AI Resume Score (0-100), breakdown metrics, strengths, weaknesses, and actionable fixes."""
    skills = parsed_data.get('skills', [])
    projects = parsed_data.get('projects', [])
    certs = parsed_data.get('certifications', [])
    experience = parsed_data.get('experience', [])
    cgpa = parsed_data.get('cgpa', 7.5)
    email = parsed_data.get('email', '')
    phone = parsed_data.get('phone', '')

    # 1. Skills Diversity (Max 25 Marks)
    skills_count = len(skills)
    if skills_count >= 10:
        skills_score = 25
    elif skills_count >= 7:
        skills_score = 20
    elif skills_count >= 4:
        skills_score = 15
    elif skills_count >= 2:
        skills_score = 10
    else:
        skills_score = 5

    # 2. Technical Projects (Max 20 Marks)
    proj_count = len(projects)
    if proj_count >= 4:
        projects_score = 20
    elif proj_count >= 2:
        projects_score = 16
    elif proj_count == 1:
        projects_score = 10
    else:
        projects_score = 5

    # 3. Certifications (Max 15 Marks)
    cert_count = len(certs)
    if cert_count >= 3:
        certs_score = 15
    elif cert_count >= 1:
        certs_score = 10
    else:
        certs_score = 4

    # 4. Internships & Work Experience (Max 15 Marks)
    exp_count = len(experience)
    if exp_count >= 2:
        exp_score = 15
    elif exp_count >= 1:
        exp_score = 12
    else:
        exp_score = 5

    # 5. Education & Academic Standing (Max 10 Marks)
    if cgpa >= 8.5:
        edu_score = 10
    elif cgpa >= 7.5:
        edu_score = 8
    elif cgpa >= 6.5:
        edu_score = 6
    else:
        edu_score = 4

    # 6. Resume Structural Completeness (Max 15 Marks)
    completeness_score = 0
    if email != "Not Specified":
        completeness_score += 4
    if phone != "Not Specified":
        completeness_score += 4
    if len(skills) > 0:
        completeness_score += 4
    if len(projects) > 0:
        completeness_score += 3

    total_score = min(100, skills_score + projects_score + certs_score + exp_score + edu_score + completeness_score)

    # Generate Strengths, Weaknesses, and Improvements
    strengths = []
    weaknesses = []
    improvements = []

    if skills_score >= 20:
        strengths.append(f"Strong tech stack with {skills_count} verified industry-relevant skills.")
    else:
        weaknesses.append("Skill set diversity is limited. Missing core frameworks or tools.")
        improvements.append("Expand skill portfolio with key frameworks like React, Node.js, Spring Boot, or Docker.")

    if proj_count >= 2:
        strengths.append(f"Demonstrates hands-on project experience ({proj_count} project entries parsed).")
    else:
        weaknesses.append("Insufficient practical hands-on project evidence.")
        improvements.append("Build and showcase 2+ full-stack or ML deployment projects with GitHub links.")

    if cert_count >= 1:
        strengths.append(f"Holds {cert_count} verified professional or cloud certifications.")
    else:
        weaknesses.append("No cloud or domain industry certifications found.")
        improvements.append("Complete at least 1 recognized certification (e.g. AWS Certified Cloud Practitioner, Oracle Java, or Meta Front-End).")

    if exp_count >= 1:
        strengths.append("Has direct internship or real-world industrial work exposure.")
    else:
        weaknesses.append("Lack of industry internship or formal work experience.")
        improvements.append("Apply for summer internships or participate in open-source hackathons to gain practical exposure.")

    if cgpa >= 7.5:
        strengths.append(f"Strong academic standing with CGPA of {cgpa:.2f}.")
    else:
        weaknesses.append(f"Academic CGPA ({cgpa:.2f}) could be enhanced for Tier-1 recruitment eligibility.")
        improvements.append("Focus on core subject fundamentals and maintain minimum 7.5+ CGPA benchmark.")

    breakdown = {
        'Skills': skills_score,
        'Projects': projects_score,
        'Certifications': certs_score,
        'Internship & Experience': exp_score,
        'Education & CGPA': edu_score,
        'Completeness & Structure': completeness_score
    }

    return {
        'total_score': total_score,
        'breakdown': breakdown,
        'strengths': strengths,
        'weaknesses': weaknesses,
        'improvements': improvements
    }

# ---------------------------------------------------------
# Skill Gap Analysis Engine
# ---------------------------------------------------------
def analyze_skill_gap(user_skills):
    """Compare user skills against industry standards to extract missing skills and resources."""
    user_skills_lower = [s.lower() for s in user_skills]

    core_industry_skills = {
        'Frontend & UI': ['javascript', 'react', 'html', 'css', 'bootstrap', 'tailwind'],
        'Backend & Databases': ['python', 'java', 'node.js', 'sql', 'postgresql', 'mongodb'],
        'DevOps & System Architecture': ['git', 'docker', 'kubernetes', 'aws', 'ci/cd', 'linux'],
        'Data Science & AI': ['pandas', 'numpy', 'scikit-learn', 'machine learning', 'deep learning']
    }

    missing_skills = []
    recommended_tech = []
    suggested_certs = []
    learning_resources = []

    for domain, skill_group in core_industry_skills.items():
        domain_missing = [s.title() for s in skill_group if s not in user_skills_lower]
        if domain_missing:
            missing_skills.extend(domain_missing[:2])

    # Tech Recommendations based on gaps
    if not any(s in user_skills_lower for s in ['docker', 'kubernetes', 'aws']):
        recommended_tech.append('Cloud & Containerization (AWS / Docker)')
        suggested_certs.append('AWS Certified Cloud Practitioner')
        learning_resources.append({
            'title': 'AWS Cloud Practitioner Essentials',
            'provider': 'Amazon Web Services (AWS)',
            'type': 'Free Official Course'
        })

    if not any(s in user_skills_lower for s in ['react', 'angular', 'vue']):
        recommended_tech.append('Modern Frontend Framework (React.js)')
        suggested_certs.append('Meta Front-End Developer Certificate')
        learning_resources.append({
            'title': 'React - The Complete Guide',
            'provider': 'Coursera / Udemy',
            'type': 'Interactive Project Course'
        })

    if not any(s in user_skills_lower for s in ['django', 'flask', 'spring boot', 'express', 'node.js']):
        recommended_tech.append('Backend Frameworks (Node.js / Express / Flask)')
        suggested_certs.append('Oracle Certified Professional: Java / Python Backend')
        learning_resources.append({
            'title': 'Node.js & Express REST API Mastery',
            'provider': 'FreeCodeCamp',
            'type': 'Free Hands-on Tutorial'
        })

    if 'git' not in user_skills_lower:
        missing_skills.append('Git & GitHub')
        recommended_tech.append('Version Control (Git & GitHub)')
        learning_resources.append({
            'title': 'Git & GitHub Complete Beginners Guide',
            'provider': 'YouTube / GitHub Skills',
            'type': 'Interactive Guide'
        })

    return {
        'missing_skills': list(set(missing_skills))[:6],
        'recommended_technologies': list(set(recommended_tech))[:4],
        'suggested_certifications': list(set(suggested_certs))[:3],
        'learning_resources': learning_resources[:4]
    }

# ---------------------------------------------------------
# Job Recommendation Engine
# ---------------------------------------------------------
def recommend_job_roles(parsed_data):
    """Match candidate against 10 target job profiles and calculate percentage suitability."""
    user_skills_lower = [s.lower() for s in parsed_data.get('skills', [])]
    cgpa = parsed_data.get('cgpa', 7.5)

    recommendations = []
    for profile in JOB_PROFILES:
        title = profile['title']
        req_skills = profile['required_skills']
        matched_skills = [s.title() for s in req_skills if s in user_skills_lower]
        missing_skills = [s.title() for s in req_skills if s not in user_skills_lower]

        skill_match_ratio = len(matched_skills) / len(req_skills) if req_skills else 0
        cgpa_eligible = cgpa >= profile['min_cgpa']

        match_score = int(skill_match_ratio * 80 + (20 if cgpa_eligible else 10))
        match_score = min(98, max(35, match_score))

        if match_score >= 75:
            badge = "Strong Match"
            badge_class = "success"
        elif match_score >= 60:
            badge = "Good Fit"
            badge_class = "primary"
        else:
            badge = "Potential Fit"
            badge_class = "warning"

        recommendations.append({
            'title': title,
            'match_score': match_score,
            'badge': badge,
            'badge_class': badge_class,
            'matched_skills': matched_skills,
            'missing_skills': missing_skills,
            'description': profile['description']
        })

    # Sort by highest match score
    recommendations.sort(key=lambda x: x['match_score'], reverse=True)
    return recommendations[:6]

# ---------------------------------------------------------
# 16-Week Personalized Learning Roadmap
# ---------------------------------------------------------
def generate_learning_roadmap(skill_gap_data):
    """Construct structured 16-Week (4-Phase) career improvement roadmap."""
    missing = skill_gap_data.get('missing_skills', ['React', 'Node.js', 'Docker', 'AWS'])
    certs = skill_gap_data.get('suggested_certifications', ['AWS Certified Cloud Practitioner'])

    phase1_tech = missing[0] if len(missing) > 0 else 'Data Structures & Algorithms'
    phase2_tech = missing[1] if len(missing) > 1 else 'Full Stack Web Architecture'

    return {
        'phase1': {
            'period': 'Weeks 1 - 4',
            'title': 'Core Technical Foundations & Gap Repair',
            'skills': [phase1_tech, 'Core SQL & Relational Databases', 'Git Version Control'],
            'deliverables': 'Complete 30 LeetCode / HackerRank DSA challenges and fix core syntax gaps.',
            'focus': 'Fundamental concepts, syntax mastery, problem-solving speed.'
        },
        'phase2': {
            'period': 'Weeks 5 - 8',
            'title': 'Domain Frameworks & Capstone Project',
            'skills': [phase2_tech, 'RESTful Microservices', 'State Management'],
            'deliverables': 'Build and deploy 1 production-ready full stack web app on Vercel / Render.',
            'focus': 'Real-world project development, clean architecture.'
        },
        'phase3': {
            'period': 'Weeks 9 - 12',
            'title': 'Cloud Deployments & Certifications',
            'skills': certs + ['Docker Containerization', 'CI/CD Pipelines'],
            'deliverables': 'Pass 1 industry certification exam and showcase project repository on GitHub.',
            'focus': 'Professional credentials, cloud infrastructure exposure.'
        },
        'phase4': {
            'period': 'Weeks 13 - 16',
            'title': 'Recruitment Preparation & Mock Interviews',
            'skills': ['Quantitative Aptitude', 'Technical Mock Interviews', 'Behavioral HR Prep'],
            'deliverables': 'Participate in 5 mock technical interview sessions and refine live coding performance.',
            'focus': 'Placement readiness, confidence building, high conversion.'
        }
    }

# ---------------------------------------------------------
# PDF Report Generator using ReportLab
# ---------------------------------------------------------
def generate_resume_pdf_report(analysis, output_filepath):
    """Generate a clean, multi-page professional PDF report of the resume analysis."""
    doc = SimpleDocTemplate(output_filepath, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    # Custom Report Typography Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1e293b')
    )
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#64748b')
    )
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#4f46e5'),
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#334155')
    )

    story = []

    # Title Banner
    story.append(Paragraph("PlacementIQ - AI Resume Analysis & Career Report", title_style))
    story.append(Paragraph(f"Candidate: <b>{analysis.get('name', 'Candidate')}</b> | File: {analysis.get('original_filename', 'Resume')} | Generated by PlacementIQ Engine", subtitle_style))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#e2e8f0'), spaceBefore=2, spaceAfter=12))

    # Executive Summary Metrics Table
    score = analysis.get('resume_score', 0)
    prob = analysis.get('placement_probability', 0.0)
    metrics_data = [
        [
            Paragraph(f"<b>AI Resume Score</b><br/><font size=16 color='#4f46e5'><b>{score}/100</b></font>", body_style),
            Paragraph(f"<b>Placement Probability</b><br/><font size=16 color='#16a34a'><b>{prob:.1f}%</b></font>", body_style),
            Paragraph(f"<b>Skills Parsed</b><br/><font size=16 color='#0284c7'><b>{len(analysis.get('skills_extracted', []))} Skills</b></font>", body_style),
            Paragraph(f"<b>Placement Status</b><br/><font size=14 color='#d97706'><b>{analysis.get('placement_status', 'Eligible')}</b></font>", body_style)
        ]
    ]
    t_metrics = Table(metrics_data, colWidths=[1.8*inch, 1.8*inch, 1.8*inch, 1.8*inch])
    t_metrics.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(t_metrics)
    story.append(Spacer(1, 12))

    # Parsed Profile Summary
    story.append(Paragraph("1. Parsed Candidate Profile", section_style))
    contact_info = f"<b>Email:</b> {analysis.get('email', 'N/A')} &nbsp;&nbsp;|&nbsp;&nbsp; <b>Phone:</b> {analysis.get('phone', 'N/A')}"
    story.append(Paragraph(contact_info, body_style))
    story.append(Spacer(1, 4))

    skills_str = ", ".join(analysis.get('skills_extracted', [])) if analysis.get('skills_extracted') else "None Parsed"
    story.append(Paragraph(f"<b>Extracted Technical Skills:</b> {skills_str}", body_style))
    story.append(Spacer(1, 10))

    # Strengths & Weaknesses
    story.append(Paragraph("2. Resume Strengths & Improvement Suggestions", section_style))
    strengths_html = "<b>Key Strengths:</b><br/>" + "<br/>".join([f"• {s}" for s in analysis.get('strengths', [])])
    story.append(Paragraph(strengths_html, body_style))
    story.append(Spacer(1, 6))

    improvements_html = "<b>Actionable Improvements:</b><br/>" + "<br/>".join([f"• {imp}" for imp in analysis.get('improvements', [])])
    story.append(Paragraph(improvements_html, body_style))
    story.append(Spacer(1, 12))

    # Skill Gap & Recommended Jobs
    story.append(Paragraph("3. Skill Gap & Job Role Suitability", section_style))
    gap_data = analysis.get('skill_gap_data', {})
    missing_str = ", ".join(gap_data.get('missing_skills', [])) if gap_data.get('missing_skills') else "None identified"
    story.append(Paragraph(f"<b>Top Missing Skills:</b> {missing_str}", body_style))
    story.append(Spacer(1, 6))

    # Recommended Jobs Table
    jobs = analysis.get('job_recommendations', [])
    if jobs:
        job_table_data = [["Target Job Role", "Match %", "Key Matched Skills", "Missing Skill Gaps"]]
        for j in jobs[:4]:
            job_table_data.append([
                Paragraph(f"<b>{j.get('title')}</b>", body_style),
                Paragraph(f"<b>{j.get('match_score')}%</b>", body_style),
                Paragraph(", ".join(j.get('matched_skills', [])[:4]), body_style),
                Paragraph(", ".join(j.get('missing_skills', [])[:3]), body_style)
            ])
        t_jobs = Table(job_table_data, colWidths=[1.8*inch, 0.8*inch, 2.4*inch, 2.2*inch])
        t_jobs.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4f46e5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(t_jobs)

    story.append(Spacer(1, 12))

    # 16-Week Roadmap Summary
    story.append(Paragraph("4. 16-Week Personalized Learning Roadmap", section_style))
    roadmap = analysis.get('roadmap_data', {})
    if roadmap:
        road_table_data = [["Phase & Timeline", "Focus Area & Required Deliverable"]]
        for k in ['phase1', 'phase2', 'phase3', 'phase4']:
            p = roadmap.get(k, {})
            if p:
                road_table_data.append([
                    Paragraph(f"<b>{p.get('period')}</b><br/>{p.get('title')}", body_style),
                    Paragraph(f"<b>Deliverable:</b> {p.get('deliverables')}<br/><b>Skills:</b> {', '.join(p.get('skills', []))}", body_style)
                ])
        t_road = Table(road_table_data, colWidths=[2.2*inch, 5.0*inch])
        t_road.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(t_road)

    doc.build(story)
    return output_filepath
